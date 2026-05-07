"""
RAG 文档加载器
支持加载 .docx, .pdf, .txt, .md 等格式的文档
"""
import os
from typing import List, Optional
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter


class DocumentLoader:
    """文档加载器"""
    
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        """
        初始化文档加载器
        
        Args:
            chunk_size: 文本块大小（字符数）
            chunk_overlap: 文本块重叠大小（字符数）
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            is_separator_regex=False,
        )
    
    def load_docx(self, file_path: str) -> List[Document]:
        """
        加载 .docx 文件
        
        Args:
            file_path: .docx 文件路径
            
        Returns:
            Document 列表
        """
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(file_path)
            full_text = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            
            # 合并文本
            text = "\n".join(full_text)
            
            # 分割文本
            chunks = self.text_splitter.split_text(text)
            
            # 创建 Document 对象
            documents = []
            for i, chunk in enumerate(chunks):
                doc = Document(
                    page_content=chunk,
                    metadata={
                        "source": file_path,
                        "file_type": "docx",
                        "chunk_id": i,
                        "title": "穿搭核心原则"
                    }
                )
                documents.append(doc)
            
            return documents
            
        except ImportError:
            # 如果没有 python-docx，使用备用方法
            return self._load_docx_fallback(file_path)
        except Exception as e:
            raise Exception(f"加载 .docx 文件失败: {str(e)}")
    
    def _load_docx_fallback(self, file_path: str) -> List[Document]:
        """备用方法：解压 docx 并提取文本"""
        import zipfile
        import xml.etree.ElementTree as ET
        
        # 解压 docx (其实是个 zip 文件)
        with zipfile.ZipFile(file_path, 'r') as zip_ref:
            xml_content = zip_ref.read('word/document.xml')
        
        # 解析 XML
        root = ET.fromstring(xml_content)
        
        # 命名空间
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        # 提取所有文本
        paragraphs = []
        for para in root.findall('.//w:p', ns):
            texts = []
            for t in para.findall('.//w:t', ns):
                if t.text:
                    texts.append(t.text)
            if texts:
                paragraphs.append(''.join(texts))
        
        # 合并文本
        full_text = "\n".join(paragraphs)
        
        # 分割文本
        chunks = self.text_splitter.split_text(full_text)
        
        # 创建 Document 对象
        documents = []
        for i, chunk in enumerate(chunks):
            doc = Document(
                page_content=chunk,
                metadata={
                    "source": file_path,
                    "file_type": "docx",
                    "chunk_id": i,
                    "title": "穿搭核心原则"
                }
            )
            documents.append(doc)
        
        return documents
    
    def load_text(self, file_path: str) -> List[Document]:
        """
        加载 .txt 或 .md 文件
        
        Args:
            file_path: 文本文件路径
            
        Returns:
            Document 列表
        """
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        # 分割文本
        chunks = self.text_splitter.split_text(text)
        
        # 创建 Document 对象
        documents = []
        for i, chunk in enumerate(chunks):
            doc = Document(
                page_content=chunk,
                metadata={
                    "source": file_path,
                    "file_type": "text",
                    "chunk_id": i
                }
            )
            documents.append(doc)
        
        return documents
    
    def load_file(self, file_path: str) -> List[Document]:
        """
        自动识别文件类型并加载
        
        Args:
            file_path: 文件路径
            
        Returns:
            Document 列表
        """
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.docx':
            return self.load_docx(file_path)
        elif ext in ['.txt', '.md']:
            return self.load_text(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {ext}")
    
    def load_directory(self, dir_path: str) -> List[Document]:
        """
        加载目录下的所有支持的文件
        
        Args:
            dir_path: 目录路径
            
        Returns:
            Document 列表
        """
        documents = []
        
        for root, dirs, files in os.walk(dir_path):
            for file in files:
                file_path = os.path.join(root, file)
                ext = os.path.splitext(file)[1].lower()
                
                if ext in ['.docx', '.txt', '.md']:
                    try:
                        docs = self.load_file(file_path)
                        documents.extend(docs)
                        print(f"✅ 已加载: {file_path}")
                    except Exception as e:
                        print(f"❌ 加载失败 {file_path}: {str(e)}")
        
        return documents
